"""產生「熱磁材料實驗驗證計畫」詳細需求 Word 文件。輸出 docs/實驗驗證計畫_詳細需求.docx"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path("docs/實驗驗證計畫_詳細需求.docx")
CJK = "WenQuanYi Zen Hei"
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TEAL = RGBColor(0x2A, 0x9D, 0x8F)

doc = Document()
# 預設字型（含東亞）
style = doc.styles["Normal"]
style.font.name = CJK
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)


def _cjk(run):
    run.font.name = CJK
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rPr.append(rf)
    rf.set(qn("w:eastAsia"), CJK)


def h(text, level=1, color=NAVY):
    p = doc.add_heading(level=level)
    r = p.add_run(text); r.font.color.rgb = color; _cjk(r)
    return p


def para(text, bold=False, size=10.5, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    _cjk(r)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; _cjk(r)
    r = p.add_run(text); _cjk(r)
    return p


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(htext); r.bold = True; _cjk(r)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val)); _cjk(r)
            r.font.size = Pt(9.5)
    doc.add_paragraph()
    return t


# ── 標題頁 ──
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("熱磁材料實驗驗證計畫"); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = NAVY; _cjk(r)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("詳細需求、規格、時程目標與分階段方案"); r.font.size = Pt(13); r.font.color.rgb = TEAL; _cjk(r)
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("alloy_engine 專案 · 2026 · 小規模先導 → 階段閘控 → 規模化擴張"); r.font.size = Pt(10); _cjk(r)
doc.add_paragraph()

# ── 0. 摘要 ──
h("0. 執行摘要", 1)
para("本計畫把熱磁發電/製冷材料引擎的『模型預測』收斂為『實測校準』。模型已用真實 NEMAD 居禮溫度"
     "（Tc R²=0.78）與 Materials Project 磁化（Br 溫度修正後 bias −0.12T）做到相對可信，但若干絕對量"
     "（工作溫度 ΔM、一階銳度 w、複合連通度 connectivity、整機 W/η/P）唯有真實樣品的磁/熱/微結構量測"
     "才能確定。策略核心是：先以最易取得的『純 Gd』做最小規模先導，打通整條『製樣→量測→回填』流程並"
     "校準儀器；通過決策門後再分期擴張到一階材料、GA 候選、複合與整機原型。每階段設 go/no-go 決策門，"
     "未過不放大投入——以最便宜的一步攔下最貴的錯誤，避免重蹈覆轍。")

# ── 1. 背景與動機 ──
h("1. 專案背景與動機", 1)
bullet("真實 Tc baseline 已達 R²=0.78 / MAE 91°C（P/Ge 擴張後 1,380 化合物），遠勝合成代理的 −0.17。", "現況：")
bullet("Br 經溫度修正後系統偏差僅 −0.12T；發電側功率密度為理想化上界（比真實原型高約 10×）。")
bullet("複合 connectivity 影響功率增益約 43%，但定性結論（複合 ×7–12、φ*≈0.3）穩健。")
bullet("以上『絕對值待收斂』的缺口，公開資料庫不含，唯有真實樣品的 M-H/DSC/SEM 能補。", "缺口：")
bullet("目標：用最小規模、可控成本，把絕對值校準，使模型成為可落地的設計工具。", "動機：")

# ── 2. 目標 ──
h("2. 目標（總體 + SMART）", 1)
table(["代號", "目標", "量化判準（SMART）"],
      [["G-A", "流程可信", "純 Gd 的 ΔM(T)/ΔS_M/Tc 與文獻 ±20% 內，且回填可復現"],
       ["G-B", "模型校準", "回填後關鍵參數（ΔM、w、connectivity）誤差 < 25%"],
       ["G-C", "端到端驗證", "GA 候選實測 Tc 誤差 < 50°C、Top-N 排序與實測一致"],
       ["G-D", "整機絕對校準", "原型 P/V、η 與文獻 TMG 原型同量級（收斂 D12 的 ~10× 落差）"]])

# ── 3. 範圍 ──
h("3. 範圍", 1)
para("在範圍內：", bold=True)
bullet("4 類最小可驗證樣品（純 Gd、La-Fe-Si(H)、Mn-Fe-P-Si、GA 候選 Fe 系）的製備/取得與磁熱量測。")
bullet("量測數據回填 alloy_engine 模型（reference_materials、composite、properties、reference_devices）。")
bullet("小型整機 TMG 原型（Phase 3）作發電側絕對校準。")
para("不在範圍（本期）：", bold=True)
bullet("大規模成分掃描、量產製程開發、商品化裝置設計——待先導與分期驗證通過後另案。")

# ── 4. 核心策略 ──
h("4. 核心策略：先導 + 決策門 + 擴張", 1)
para("採階段閘控（stage-gate）。每階段先小規模驗證『流程＋模型量級』，通過決策門才放大投入；任一門未過，"
     "回上一階段診斷是量測問題（校準/退火/氧化）或模型問題（公式/單位/假設），修正後重做，不貿然擴張。")
table(["決策門", "時點", "通過條件", "未過處置"],
      [["G0", "Phase 0 後", "Gd 量測與文獻 ±20%、回填可復現", "修儀器/流程，重做先導"],
       ["G1", "Phase 2 後", "GA 候選端到端 sim-to-real 收斂", "回填修模型，再評估是否進原型"]])

# ── 5. 樣品需求規格 ──
h("5. 樣品需求規格", 1)
table(["#", "材料", "成分（目標）", "純度/相要求", "量", "取得方式"],
      [["M1", "純 Gd", "Gd 99.9%", "金屬 Gd、無氧化層", "~數 g", "商購（高純稀土）"],
       ["M2", "La(Fe,Si)13(H)", "La~7%, Fe~80%, Si~13%", "1:13 相（XRD 確認）", "~數 g/成分", "委外/學界合作合成"],
       ["M3", "(Mn,Fe)2(P,Si)", "Mn≈Fe≈1/3, (P+Si)≈1/3", "六方 Fe2P 型相", "~數 g/成分", "委外/學界合作合成"],
       ["M4", "GA 候選 Fe 系", "由 run_search 產出", "單相/可鑄", "~數 g×1–2", "電弧熔煉自製"]])
para("備註：M2/M3 為一階材料，需正確退火與相形成；脆性高，製樣需小心（與模型 D9 脆性懲罰一致）。"
     "故優先委外或與有經驗的實驗室合作，降低製程風險。", italic=True)

# ── 6. 樣品取得方案 ──
h("6. 樣品取得方案", 1)
table(["方式", "適用材料", "優點", "風險 / 成本", "前置"],
      [["商購", "純 Gd、純元素", "快、純度保證、低風險", "成分受限", "供應商詢價"],
       ["自製（電弧熔煉）", "Fe 系 GA 候選", "成分自由、快迭代", "需熔煉+退火設備", "盤點熔煉爐"],
       ["委外/合作", "La-Fe-Si、Mn-Fe-P", "拿到難製一階相", "週期長、需協調", "聯繫合作實驗室"]])
para("採購清單（Phase 0 啟動所需，最小）：高純 Gd 金屬、切割/封裝耗材、惰性氣氛保存。", bold=True)

# ── 7. 量測需求 ──
h("7. 量測需求（設備 × 測項 × 回填）", 1)
table(["設備", "測項", "關鍵參數", "輸出", "回填模型位置"],
      [["EDS / XRD", "成分 / 相", "相純度", "成分、相確認", "—（前置確認）"],
       ["VSM / PPMS", "M-H 迴線 / M(T)", "多溫度、±μ0H", "ΔJ、磁滯面積、w", "ΔM 估計、磁滯損耗、D5 w"],
       ["DSC", "ΔS_M / Cp / Tc", "過 Tc 掃描", "熱量、相變峰", "reference_materials、properties"],
       ["雷射閃光", "熱擴散 α", "室溫～工作溫度", "κ", "generator_design 頻率"],
       ["SEM", "微結構 / 連通度", "相分布、φ", "connectivity、φ", "composite.py（D6 最敏感）"],
       ["自製 TMG 台", "W / η / P", "磁路+熱循環", "整機效能", "reference_devices（D12 錨點）"]])

# ── 8. 階段計畫 ──
h("8. 分階段計畫", 1)
for title_, body in [
    ("Phase 0 — 先導：純 Gd（最小可行，~2 月、1 樣）",
     [("目的", "打通『製樣→量測→回填→比對文獻』整鏈，並用答案已知的 Gd 校準儀器。"),
      ("輸入", "商購高純 Gd。"),
      ("活動", "切樣→EDS/XRD→VSM M-H（多溫度）→DSC（過 Tc）→M(T) 擬合→比對文獻→回填。"),
      ("輸出", "ΔJ(T)、ΔS_M、Tc；流程可信度結論。"),
      ("決策門", "G0：與文獻 ±20%、可復現。")]),
    ("Phase 1 — 一階材料：La-Fe-Si / Mn-Fe-P（~3 月、2–3 樣）",
     [("目的", "校 D5 銳度 w、ΔS_M；驗證 D8 氫化 Tc 上修。"),
      ("輸入", "委外/合作合成的一階樣品。"),
      ("活動", "XRD 確認相→VSM/DSC→M(T) 銳度擬合 w；氫化前後各量一次比較 Tc 位移。"),
      ("輸出", "w、ΔS_M、氫化 Tc 位移；回填 reference_materials 與 D5/D8。"),
      ("決策門", "—（併入 G1 評估）。")]),
    ("Phase 2 — GA 候選 + 複合 connectivity（~3 月、~6 樣）",
     [("目的", "端到端驗證 GA 預測；校 D6 connectivity 絕對值。"),
      ("輸入", "bundle_real_tc 跑 run_search 的 Top-N；高κ基底+一階相複合。"),
      ("活動", "電弧熔煉 GA 候選→量 Tc/Br/σy 比對；複合做 SEM 量連通度→回填。"),
      ("輸出", "sim-to-real 誤差；connectivity 實測值。"),
      ("決策門", "G1：端到端收斂、排序正確。")]),
    ("Phase 3 — 整機原型（~4 月，前三階段全綠才啟動）",
     [("目的", "發電側絕對校準（收斂 D12 的 ~10× 落差）。"),
      ("輸入", "校準後最佳材料（一階+複合）。"),
      ("活動", "設計磁路/換熱/線圈→組裝→量 W/η/P-V→新增 reference_devices『本工作』錨點。"),
      ("輸出", "整機絕對效能；模型升級為絕對預測。"),
      ("決策門", "計畫收尾 / 續期評估。")]),
]:
    h(title_, 2)
    for k, v in body:
        bullet(v, bold_prefix=f"{k}：")

# ── 9. 時程與里程碑 ──
h("9. 時程與里程碑", 1)
table(["月", "里程碑", "決策門"],
      [["M2", "先導完成（Gd 流程打通）", "G0 go/no-go"],
       ["M5", "一階材料 + 氫化驗證", "—"],
       ["M8", "GA 候選 + 複合 connectivity", "G1 go/no-go"],
       ["M12", "整機原型、發電側絕對校準", "收尾/續期"]])

# ── 10. 預算與資源 ──
h("10. 預算與資源（相對量級）", 1)
table(["階段", "樣品/合成", "量測", "相對量級"],
      [["Phase 0", "商購 Gd（低）", "VSM+DSC（低）", "★ 最低"],
       ["Phase 1", "委外一階（中）", "+氫化（中）", "★★"],
       ["Phase 2", "自製+複合（中）", "+SEM+κ（中）", "★★★"],
       ["Phase 3", "原型製作（高）", "整機量測（高）", "★★★★"]])
para("精確金額依在地共儀/委外報價而定；本表供排序與分期經費申請。先用共儀與委外壓低固定成本，"
     "驗證可行再考慮自建設備。", italic=True)

# ── 11. 風險登錄 ──
h("11. 風險登錄與緩解", 1)
table(["風險", "影響", "緩解措施"],
      [["一階相未正確形成", "量到的不是目標相", "委外/合作 + XRD 先確認相"],
       ["稀土氧化 / 氫脆", "樣品劣化、數據失真", "惰性氣氛、鍍層、儘速量測"],
       ["量測未校準", "絕對值不可信", "Phase 0 用 Gd 先校準儀器"],
       ["範圍蔓延", "成本失控", "決策門閘控、分期投入"],
       ["GA 候選不可製造", "白做樣品", "先過 D9 脆性/可製造性篩選"]])

# ── 12. 數據與品質管理 ──
h("12. 數據與品質管理", 1)
bullet("每階段原始數據 + 回填參數 + 復現腳本入 git，與 alloy_engine 現有腳本銜接。")
bullet("量測前以標準樣校準；每材料至少重複量測確認再現性。")
bullet("回填後跑既有測試套件（258 項）確認模型不退步（CI 守護）。")

# ── 13. 成功判準與 KPI ──
h("13. 成功判準與 KPI", 1)
bullet("先導：Gd ΔM/ΔS/Tc 與文獻 ±20%（最關鍵，沒過先停）。")
bullet("校準：關鍵參數誤差 < 25%；相對排序不變。")
bullet("端到端：GA 候選 Tc 誤差 < 50°C、Top-N 排序一致。")
bullet("整機：原型 P/V、η 與文獻原型同量級。")

# ── 14. 角色與責任 ──
h("14. 角色與責任（最小編制）", 1)
table(["角色", "職責", "階段"],
      [["材料製備", "合成/退火/委外協調", "全程"],
       ["磁/熱量測", "VSM/DSC/κ + 數據", "全程"],
       ["建模回填", "實測→模型校準→復現", "全程"],
       ["整機工程", "原型設計與量測", "Phase 3"]])
para("可由 2–3 人 + 共儀起步；建模回填角色與既有 alloy_engine 腳本直接銜接。", italic=True)

# ── 附錄 ──
h("附錄 A：量到什麼 → 回填模型哪裡（閉環對應表）", 1)
table(["量測值", "回填位置", "對應缺陷"],
      [["工作溫度 ΔJ", "properties.magnetic_thermodynamic_score", "ΔM"],
       ["logistic 寬度 w", "magnetic_thermodynamic_score(transition_width_K)", "D5"],
       ["複合 connectivity、φ", "composite.composite_properties", "D6"],
       ["κ、Cp、ΔS_M", "reference_materials / design_tmg", "—"],
       ["氫化前後 Tc", "synthetic.hydrogenation_tc_shift_K", "D8"],
       ["磁滯損耗", "magnetocaloric_refrigeration 懲罰", "—"],
       ["整機 W/η/P", "reference_devices（新增錨點）", "D12"]])

h("附錄 B：立即下一步（啟動清單）", 1)
bullet("盤點可用共儀（XRD/SEM/DSC/VSM）與熔煉/退火/氫化設備。")
bullet("採購高純 Gd 啟動 Phase 0。")
bullet("跑 run_search.py（用 bundle_real_tc）產出 GA 候選清單，供 Phase 2 備料。")
bullet("聯繫具 La-Fe-Si / Mn-Fe-P 經驗的實驗室洽談合作合成。")

doc.save(str(OUT))
print("saved", OUT)
